from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = "/mnt/storage2/zyc/CIM比赛/赛道4/赛题2"
DOC_DIR = os.path.join(BASE_DIR, "文档")

def create_docx():
    output_path = os.path.join(BASE_DIR, '技术报告.docx')
    doc = Document()

    title = doc.add_heading('Analog CIM Design Based on FET Devices', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('赛道四：存算设计 - 赛题二')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('团队 / Team: HDUer')
    doc.add_paragraph('成员 / Member: 1')
    doc.add_paragraph('日期 / Date: 2026-05-11')

    doc.add_page_break()

    chapters = [
        ('01_器件原理.md', 'Chapter 1: FeFET Device Principles', '第一章 FeFET器件的工作原理及特性介绍'),
        ('02_工作区间.md', 'Chapter 2: Work Region Analysis', '第二章 工作区间选择的分析'),
        ('03_正负信号处理.md', 'Chapter 3: Sign Handling', '第三章 正负信号处理方案'),
        ('04_线性度分析.md', 'Chapter 4: Linearity Analysis', '第四章 线性度分析与补偿'),
        ('05_精度分析.md', 'Chapter 5: Precision Analysis', '第五章 多比特精度分析与评估'),
        ('06_总结讨论.md', 'Chapter 6: Summary', '第六章 总结与讨论'),
        ('参考文献.md', 'References', '参考文献'),
    ]

    for filename, title_en, title_cn in chapters:
        filepath = os.path.join(DOC_DIR, filename)
        if not os.path.exists(filepath):
            continue

        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        h = doc.add_heading(title_en, level=1)
        h = doc.add_heading(title_cn, level=2)

        lines = content.split('\n')
        table_data = []

        for line in lines:
            line = line.rstrip()

            if line.startswith('```') or line.startswith('~~~'):
                continue

            if line.startswith('|'):
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if all(parts):
                    table_data.append(parts)
                continue
            elif table_data:
                if len(table_data) > 1:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_data):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data):
                            cell = row.cells[j]
                            cell.text = cell_text
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    doc.add_paragraph()
                table_data = []

            if line.startswith('# '):
                continue
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=3)
            elif line.startswith('### '):
                p = doc.add_paragraph(line[4:])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                doc.add_paragraph(line, style='List Number')
            elif line.strip() and not line.startswith('---') and not line.startswith('*'):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    doc.save(output_path)
    print(f"DOCX saved: {output_path}")
    return output_path

if __name__ == "__main__":
    create_docx()
