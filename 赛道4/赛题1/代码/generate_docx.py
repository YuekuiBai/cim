from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = "/mnt/storage2/zyc/CIM比赛/赛道4/赛题1"
DOC_DIR = os.path.join(BASE_DIR, "文档")

def set_font(run, font_name='Times New Roman', east_asia='宋体', size=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)
    if size:
        run.font.size = Pt(size)
    run.bold = bold

def create_document():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_heading('神经网络在模拟存内计算系统的映射', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, east_asia='宋体', size=22, bold=True)

    subtitle = doc.add_paragraph('赛道四：存算设计 - 赛题一')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_font(run, east_asia='宋体', size=14)

    for text in ['团队: HDUer', '成员: 1人', '日期: 2026-05-10']:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_font(run, east_asia='宋体', size=12)

    doc.add_page_break()

    chapters = [
        ('01_器件原理.md', '第一章 存储器件的工作原理及特性介绍'),
        ('02_神经网络模型.md', '第二章 神经网络模型的介绍'),
        ('03_存算阵列规格.md', '第三章 存内计算阵列规格信息'),
        ('04_网络映射方案.md', '第四章 神经网络映射方案'),
        ('05_性能评估.md', '第五章 性能及效率评估'),
        ('06_总结讨论.md', '第六章 总结与讨论'),
        ('参考文献.md', '参考文献'),
    ]

    for filename, title in chapters:
        filepath = os.path.join(DOC_DIR, filename)
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()

            h = doc.add_heading(title, 1)
            for run in h.runs:
                set_font(run, east_asia='宋体', size=16, bold=True)

            lines = content.split('\n')
            table_rows = []

            for line in lines:
                line = line.rstrip()

                if line.startswith('```') or line.startswith('~~~'):
                    continue

                if line.startswith('|'):
                    parts = [p.strip() for p in line.split('|')[1:-1]]
                    if all(parts):
                        table_rows.append(parts)
                    continue
                elif table_rows:
                    if len(table_rows) > 1:
                        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        for i, row_data in enumerate(table_rows):
                            for j, cell_data in enumerate(row_data):
                                cell = table.rows[i].cells[j]
                                cell.text = cell_data
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        set_font(run, east_asia='宋体', size=10, bold=(i==0))
                    table_rows = []

                if line.startswith('# '):
                    h = doc.add_heading(line[2:], 2)
                    for run in h.runs:
                        set_font(run, east_asia='宋体', size=14, bold=True)
                elif line.startswith('## '):
                    h = doc.add_heading(line[3:], 3)
                    for run in h.runs:
                        set_font(run, east_asia='宋体', size=13, bold=True)
                elif line.startswith('### '):
                    p = doc.add_paragraph(line[4:])
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        set_font(run, east_asia='宋体', size=11)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('**', ''))
                    set_font(run, east_asia='宋体', size=11, bold=True)
                elif line.startswith('- '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    for run in p.runs:
                        set_font(run, east_asia='宋体', size=11)
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    p = doc.add_paragraph(line, style='List Number')
                    for run in p.runs:
                        set_font(run, east_asia='宋体', size=11)
                elif line.strip() and not line.startswith('---'):
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        set_font(run, east_asia='宋体', size=11)

            doc.add_page_break()

    output_path = os.path.join(BASE_DIR, '技术报告.docx')
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
